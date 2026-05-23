"""
text_to_docx.py — shared markdown-to-docx renderer.

Converts a markdown string or file to a styled .docx using python-docx.
Used by cover-letter and interview-prep (neither needs Jinja templating;
they need plain text wrapped cleanly in Word format).

Supported markdown:
  ## Heading      → Heading 2
  ### Heading     → Heading 3
  **bold text**   → bold run (inline, within a paragraph)
  - bullet item   → List Bullet style
  blank line      → paragraph break

CLI:
  python text_to_docx.py <input.md> <output.docx>

Library:
  from shared.scripts.text_to_docx import markdown_to_docx
  markdown_to_docx(text, "output.docx")
"""

import os
import re
import sys

from docx import Document
from docx.shared import Pt


def _add_paragraph_with_inline_bold(doc, style, text):
    """Add a paragraph that may contain **bold** spans."""
    para = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)
    return para


def markdown_to_docx(text, output_path):
    """Convert a markdown string to a .docx file at output_path."""
    doc = Document()

    # Remove default empty paragraph Word adds
    for element in doc.paragraphs:
        element._element.getparent().remove(element._element)

    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- ") or line.startswith("* "):
            _add_paragraph_with_inline_bold(doc, "List Bullet", line[2:].strip())
        elif line.strip() == "":
            pass  # paragraph break; next non-empty line starts a new paragraph
        else:
            _add_paragraph_with_inline_bold(doc, "Normal", line.strip())

        i += 1

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python text_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    input_path, output_path = sys.argv[1], sys.argv[2]
    if not os.path.isfile(input_path):
        print(f"Error: file not found: {input_path}", file=sys.stderr)
        sys.exit(1)
    with open(input_path, encoding="utf-8") as f:
        source = f.read()
    markdown_to_docx(source, output_path)
    print(f"Saved: {output_path}")
